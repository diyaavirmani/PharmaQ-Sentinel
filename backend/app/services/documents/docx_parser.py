from __future__ import annotations

from pathlib import Path
from typing import ClassVar

from docx import Document

from app.core.exceptions import PharmaQSentinelError
from app.services.documents.base import DocumentParser
from app.services.documents.schemas import DocumentSegment, DocumentTextResult


class DocxDocumentParser(DocumentParser):
    document_type: ClassVar[str] = "DOCX"
    supported_mime_types: ClassVar[set[str]] = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }

    def parse(self, path: Path) -> DocumentTextResult:
        try:
            document = Document(path)
        except Exception as exc:
            raise PharmaQSentinelError("Malformed DOCX could not be parsed", status_code=422) from exc

        segments: list[DocumentSegment] = []
        text_parts: list[str] = []
        offset = 0
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = " ".join(paragraph.text.split())
            if not text:
                continue
            start = offset
            end = start + len(text)
            segments.append(
                DocumentSegment(
                    segment_id=f"docx-p{index}",
                    page_number=None,
                    paragraph_index=index,
                    text=text,
                    start_offset=start,
                    end_offset=end,
                )
            )
            text_parts.append(text)
            offset = end + 1

        return DocumentTextResult(
            document_type=self.document_type,
            detected_mime_type=next(iter(self.supported_mime_types)),
            text="\n".join(text_parts),
            segments=segments,
            metadata={"paragraph_count": len(document.paragraphs)},
        )
