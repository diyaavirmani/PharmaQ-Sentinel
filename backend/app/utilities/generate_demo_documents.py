from __future__ import annotations

import json
from email.message import EmailMessage
from pathlib import Path

import fitz
from docx import Document

DEMO_NOTICE = "DEMONSTRATION DATA - NOT A REAL PHARMACEUTICAL RECORD"


def _output_directory() -> Path:
    path = Path(__file__).resolve().parents[2] / "storage" / "demo_documents"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _write_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    text = (
        f"{DEMO_NOTICE}\n\n"
        "Apollo Pharmacy reported 12 discoloured Amoxicillin Capsules 500 mg from batch BMX240602. "
        "Manufacturing date March 2026 and expiry date February 2028. Sample is available."
    )
    page.insert_textbox((72, 72, 520, 260), text, fontsize=11)
    document.save(path)
    document.close()


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph(DEMO_NOTICE)
    document.add_paragraph(
        "Demo QA contact reported an API assay discrepancy for Amoxicillin API batch BMX240603. "
        "The reported result is below expected internal specification and requires QA review."
    )
    document.save(path)


def _write_txt(path: Path) -> None:
    path.write_text(
        f"{DEMO_NOTICE}\n\n"
        "Jaipur distributor reported blister leakage involving 18 strips of Paracetamol Tablets 500 mg "
        "from batch BMX240604.",
        encoding="utf-8",
    )


def _write_eml(path: Path) -> None:
    message = EmailMessage()
    message["Subject"] = "Demo customer complaint"
    message["From"] = "demo.customer@example.test"
    message["To"] = "qa@example.test"
    message.set_content(
        f"{DEMO_NOTICE}\n\n"
        "Mumbai demo customer reported packaging leakage for Amoxicillin Capsules 500 mg from batch BMX240602."
    )
    path.write_text(message.as_string(), encoding="utf-8")


def main() -> None:
    output = _output_directory()
    files = {
        "amoxicillin_capsule_discolouration_pdf": output / "amoxicillin_capsule_discolouration.pdf",
        "api_assay_complaint_docx": output / "api_assay_complaint.docx",
        "packaging_leakage_txt": output / "packaging_leakage_complaint.txt",
        "customer_complaint_eml": output / "customer_complaint.eml",
    }
    _write_pdf(files["amoxicillin_capsule_discolouration_pdf"])
    _write_docx(files["api_assay_complaint_docx"])
    _write_txt(files["packaging_leakage_txt"])
    _write_eml(files["customer_complaint_eml"])
    print(json.dumps({key: str(path) for key, path in files.items()}, indent=2))


if __name__ == "__main__":
    main()
